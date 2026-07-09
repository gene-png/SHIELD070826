"""B-4 (full gap list in XLSX; "Top N of <total>" narrative) and B-5 (ZT
roadmap + DOCX Answers) exporter regressions.

These assert on the CONTENT of the generated files (openpyxl / python-docx),
not just that bytes were produced -- the target-mismatch class of defect
survived a green suite precisely because nobody parsed the artifacts.
"""

from __future__ import annotations

import io
import uuid

import pytest
from app.csf.catalog import SUBCATEGORIES
from app.csf.exporters import build_context as build_csf_context
from app.csf.exporters import render_docx as render_csf_docx
from app.csf.exporters import render_xlsx as render_csf_xlsx
from app.csf.gap import analyze as analyze_csf_gaps
from app.csf.scoring import compute as compute_csf_score
from app.models.csf_assessment import CsfAnswer, CsfAssessment, CsfAssessmentStatus
from app.models.zt_assessment import (
    ZtAnswer,
    ZtAssessment,
    ZtAssessmentStatus,
    ZtFramework,
)
from app.zt.catalog import capabilities
from app.zt.exporters import build_context as build_zt_context
from app.zt.exporters import render_docx as render_zt_docx
from app.zt.exporters import render_xlsx as render_zt_xlsx
from app.zt.maturity import ZtFrameworkCode
from app.zt.scoring import analyze_gaps as analyze_zt_gaps
from app.zt.scoring import build_roadmap
from app.zt.scoring import compute as compute_zt_score
from docx import Document
from openpyxl import load_workbook


# ---------------------------------------------------------------------------
# ZT fixtures
# ---------------------------------------------------------------------------
def _zt_ctx(*, stage: int = 3, target: int = 4):
    fw = ZtFrameworkCode.CISA_ZTMM_2_0
    a = ZtAssessment(
        id=uuid.uuid4(),
        service_id=uuid.uuid4(),
        framework=ZtFramework.CISA_ZTMM_2_0,
        version=1,
        status=ZtAssessmentStatus.APPROVED,
    )
    answers = [
        ZtAnswer(
            id=uuid.uuid4(),
            assessment_id=a.id,
            capability_code=c.code,
            maturity_stage=stage,
        )
        for c in capabilities(fw)
    ]
    stage_map = {ans.capability_code: ans.maturity_stage for ans in answers}
    score = compute_zt_score(fw, stage_map)
    # top_n=None mirrors the finalize path: the XLSX must carry the FULL list.
    gap = analyze_zt_gaps(fw, stage_map, target_stage=target, top_n=None)
    roadmap = build_roadmap(gap.gaps)
    ctx = build_zt_context(
        client_legal_name="Atlas Defense Solutions",
        service_title="Zero Trust Assessment",
        framework=fw,
        assessment=a,
        answers=answers,
        score=score,
        gap=gap,
        roadmap=roadmap,
    )
    return ctx, gap, roadmap


# ---------------------------------------------------------------------------
# B-4: XLSX carries every gap; PDF/DOCX narrative says "Top N of <total>"
# ---------------------------------------------------------------------------
@pytest.mark.unit
def test_zt_xlsx_gap_plan_lists_full_list_not_capped_at_20() -> None:
    ctx, gap, _ = _zt_ctx()
    assert gap.total_gap_count > 20, "need 25+ gaps to prove truncation is gone"
    wb = load_workbook(io.BytesIO(render_zt_xlsx(ctx)))
    ws = wb["Gap Plan"]
    # Header + one row per gap; equals the TRUE total, not 20.
    assert ws.max_row == gap.total_gap_count + 1


@pytest.mark.unit
def test_zt_docx_gap_heading_states_top_20_of_total() -> None:
    ctx, gap, _ = _zt_ctx()
    doc = Document(io.BytesIO(render_zt_docx(ctx)))
    headings = [p.text for p in doc.paragraphs]
    assert any(
        t.startswith("Top ") and f"of {gap.total_gap_count}" in t for t in headings
    ), headings


@pytest.mark.unit
def test_csf_xlsx_gap_plan_lists_full_list_not_capped_at_20() -> None:
    a = CsfAssessment(
        id=uuid.uuid4(),
        service_id=uuid.uuid4(),
        version=1,
        status=CsfAssessmentStatus.APPROVED,
    )
    answers = [
        CsfAnswer(
            id=uuid.uuid4(),
            assessment_id=a.id,
            subcategory_code=sc.code,
            maturity_tier=3,
        )
        for sc in SUBCATEGORIES
    ]
    tier_map = {ans.subcategory_code: ans.maturity_tier for ans in answers}
    score = compute_csf_score(tier_map)
    gap = analyze_csf_gaps(tier_map, target_tier=4, top_n=None)
    assert gap.total_gap_count > 20
    ctx = build_csf_context(
        client_legal_name="Atlas Defense Solutions",
        service_title="NIST CSF 2.0 Assessment",
        assessment=a,
        answers=answers,
        score=score,
        gap=gap,
    )
    wb = load_workbook(io.BytesIO(render_csf_xlsx(ctx)))
    ws = wb["Gap Plan"]
    assert ws.max_row == gap.total_gap_count + 1

    doc = Document(io.BytesIO(render_csf_docx(ctx)))
    headings = [p.text for p in doc.paragraphs]
    assert any(
        t.startswith("Top ") and f"of {gap.total_gap_count}" in t for t in headings
    ), headings


# ---------------------------------------------------------------------------
# B-5: ZT roadmap sheet + section, and DOCX Answers section
# ---------------------------------------------------------------------------
@pytest.mark.unit
def test_zt_xlsx_roadmap_sheet_matches_build_roadmap() -> None:
    ctx, _, roadmap = _zt_ctx()
    wb = load_workbook(io.BytesIO(render_zt_xlsx(ctx)))
    assert "Roadmap" in wb.sheetnames
    ws = wb["Roadmap"]
    # Header + one row per roadmap item.
    assert ws.max_row == len(roadmap) + 1
    assert [c.value for c in ws[1]] == ["Month", "Capability", "Pillar", "From stage", "To stage"]


@pytest.mark.unit
def test_zt_docx_has_answers_section() -> None:
    ctx, _, _ = _zt_ctx()
    doc = Document(io.BytesIO(render_zt_docx(ctx)))
    assert any(p.text == "Answers" for p in doc.paragraphs), [p.text for p in doc.paragraphs]


@pytest.mark.unit
def test_zt_docx_has_roadmap_section() -> None:
    ctx, _, _ = _zt_ctx()
    doc = Document(io.BytesIO(render_zt_docx(ctx)))
    assert any(p.text == "Remediation roadmap" for p in doc.paragraphs)
