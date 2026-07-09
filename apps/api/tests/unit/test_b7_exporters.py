"""B-7 exporter regressions:

  (1) ATT&CK DOCX/PDF remediation list is priority-ordered (weakest tactic
      coverage first, gaps before partials), not an alphabetical GAP-only dump.
  (2) Risk DOCX carries the 5x5 Likelihood x Impact matrix the PDF renders.
  (3) ATT&CK DOCX falls back to the technique code (not a blank cell) on a
      catalog KeyError, matching the PDF.

All assertions parse the generated file content.
"""

from __future__ import annotations

import io
from types import SimpleNamespace

import pytest
from app.attack.analytics import CoverageRollup, TacticCoverage
from app.attack.coverage import CoverageStatus
from app.attack.exporters import build_context as build_attack_context
from app.attack.exporters import render_docx as render_attack_docx
from app.risk.exporters import build_context as build_risk_context
from app.risk.exporters import render_docx as render_risk_docx
from docx import Document


# ---------------------------------------------------------------------------
# ATT&CK helpers
# ---------------------------------------------------------------------------
def _cov(code: str, status: str):
    return SimpleNamespace(technique_code=code, status=status, notes=None)


def _rollup(tactic_pcts: dict[str, float]) -> CoverageRollup:
    by_tactic = tuple(
        TacticCoverage(
            tactic_id=tid,
            tactic_name=tid,
            technique_count=1,
            sub_technique_count=0,
            covered=0,
            partial=0,
            gap=1,
            not_applicable=0,
            unscored=0,
            coverage_pct=pct,
        )
        for tid, pct in tactic_pcts.items()
    )
    return CoverageRollup(
        total_techniques=1,
        total_sub_techniques=0,
        scored_count=1,
        unscored_count=0,
        covered=0,
        partial=0,
        gap=1,
        not_applicable=0,
        coverage_pct=0.0,
        by_tactic=by_tactic,
    )


def _attack_ctx(coverage, rollup):
    return build_attack_context(
        client_legal_name="Atlas Defense Solutions",
        service_title="MITRE ATT&CK Coverage",
        assessment=SimpleNamespace(version=1),
        coverage=coverage,
        rollup=rollup,
    )


def _remediation_table(doc):
    return next(t for t in doc.tables if t.rows[0].cells[0].text == "Code")


@pytest.mark.unit
def test_attack_docx_remediation_sorted_by_tactic_coverage_then_gap_before_partial() -> None:
    # T1189 -> tactic TA0001 (coverage 90%, well-covered tactic)
    # T1592, T1591 -> tactic TA0043 (coverage 10%, weak tactic)
    # Alphabetical (the bug) would put T1189 first and drop the PARTIAL entirely.
    # Priority order puts the weak-tactic entries first, gap (T1592) before the
    # partial (T1591), and the well-covered T1189 last.
    coverage = [
        _cov("T1189", CoverageStatus.GAP.value),
        _cov("T1592", CoverageStatus.GAP.value),
        _cov("T1591", CoverageStatus.PARTIAL.value),
    ]
    ctx = _attack_ctx(coverage, _rollup({"TA0001": 90.0, "TA0043": 10.0}))
    doc = Document(io.BytesIO(render_attack_docx(ctx)))
    table = _remediation_table(doc)
    codes = [r.cells[0].text for r in table.rows[1:]]
    assert codes == ["T1592", "T1591", "T1189"], codes


@pytest.mark.unit
def test_attack_docx_section_title_states_the_sort_rule() -> None:
    ctx = _attack_ctx([_cov("T1592", CoverageStatus.GAP.value)], _rollup({"TA0043": 10.0}))
    doc = Document(io.BytesIO(render_attack_docx(ctx)))
    text = " ".join(p.text for p in doc.paragraphs).lower()
    assert "weakest tactic coverage first" in text
    assert "gaps before partials" in text


@pytest.mark.unit
def test_attack_docx_unknown_technique_falls_back_to_code_not_blank() -> None:
    ctx = _attack_ctx([_cov("T9999", CoverageStatus.GAP.value)], _rollup({"TA0001": 50.0}))
    doc = Document(io.BytesIO(render_attack_docx(ctx)))
    table = _remediation_table(doc)
    row = table.rows[1]
    assert row.cells[0].text == "T9999"
    # The name column falls back to the code, never a blank cell (matches the PDF).
    assert row.cells[1].text == "T9999"


# ---------------------------------------------------------------------------
# Risk DOCX 5x5 matrix
# ---------------------------------------------------------------------------
def _risk_entry(likelihood="high", impact="major"):
    return SimpleNamespace(
        title="Weak MFA on admin consoles",
        description="desc",
        axis="detection",
        source="attack",
        source_id="T1078",
        linked_techniques=["T1078"],
        linked_controls=[],
        likelihood=likelihood,
        impact=impact,
        tier="high",
        compensating_controls="",
        residual_risk="",
        recommended_action="remediate",
        rationale="",
        origin="ai",
        trust="",
    )


@pytest.mark.unit
def test_risk_docx_carries_the_5x5_matrix() -> None:
    ctx = build_risk_context(
        client_legal_name="Atlas Defense Solutions",
        version=1,
        entries=[_risk_entry(), _risk_entry(likelihood="low", impact="minor")],
    )
    doc = Document(io.BytesIO(render_risk_docx(ctx)))
    assert any(p.text == "Likelihood x Impact matrix" for p in doc.paragraphs)

    impacts = ["Negligible", "Minor", "Moderate", "Major", "Catastrophic"]
    matrix = None
    for t in doc.tables:
        header = [c.text for c in t.rows[0].cells]
        if header[1:] == impacts:
            matrix = t
            break
    assert matrix is not None, [[c.text for c in t.rows[0].cells] for t in doc.tables]
    # 1 header row + 5 likelihood rows; 1 label col + 5 impact cols.
    assert len(matrix.rows) == 6
    assert len(matrix.columns) == 6
